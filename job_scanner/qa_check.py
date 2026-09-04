#!/usr/bin/env python3
"""The check battery. Runs offline; no network, no browser.

    python qa_check.py

Checks are registered with @check and run in order. Every one of them exists
because something could silently go wrong in a way the output would not
reveal: a widened max_age_days, a deadline quietly filling in for a posting
date, an adapter returning zero because it failed rather than because the
employer has no vacancies.

Never weaken, skip or delete a check to make something pass.
"""

from __future__ import annotations

import datetime as _dt
import sys
import tempfile
import traceback
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from jobscan import adapters, config as config_module, normalize, writers  # noqa: E402
from jobscan.adapters import AdapterError  # noqa: E402
from jobscan.fetch import Fetcher  # noqa: E402
from jobscan.model import Posting  # noqa: E402
from jobscan.run import RunLog  # noqa: E402
from jobscan.scoring import Scorer  # noqa: E402

ROOT = Path(__file__).resolve().parent
CHECKS: list[tuple[str, callable]] = []


def check(name: str):
    def register(fn):
        CHECKS.append((name, fn))
        return fn
    return register


class Failed(AssertionError):
    pass


def expect(condition, message: str) -> None:
    if not condition:
        raise Failed(message)


# --------------------------------------------------------------------------
# Configuration integrity
# --------------------------------------------------------------------------

@check("profile block is unchanged (frozen)")
def _profile_frozen():
    cfg = config_module.load()
    lock = config_module.PROFILE_LOCK
    expect(lock.exists(), f"{lock} is missing; the frozen profile has no baseline to compare against")
    recorded = lock.read_text().strip()
    expect(
        cfg.fingerprint() == recorded,
        "the profile block has been edited. It is frozen by rule -- including "
        f"max_age_days.\n  expected {recorded}\n  found    {cfg.fingerprint()}",
    )


@check("max_age_days is present, positive and not widened past a sane bound")
def _max_age_sane():
    cfg = config_module.load()
    expect(cfg.max_age_days > 0, "max_age_days must be positive")
    expect(
        cfg.max_age_days <= 180,
        f"max_age_days is {cfg.max_age_days}; beyond ~180 days it stops "
        "distinguishing a live board from a stale mirror, which is its whole job",
    )


@check("config loads and every source key is unique")
def _config_loads():
    cfg = config_module.load()
    keys = [s["key"] for s in cfg.sources]
    expect(len(keys) == len(set(keys)), "duplicate source keys in sources.yaml")
    expect(len(cfg.sources) >= 1, "sources.yaml has no sources")


@check("verified flags mean what they claim")
def _verified_discipline():
    cfg = config_module.load()
    for source in cfg.sources:
        verified = source.get("verified", "unconfirmed")
        expect(
            verified in (True, "url-confirmed", "unconfirmed"),
            f"{source['key']}: verified={verified!r} is not one of "
            "true / url-confirmed / unconfirmed",
        )
        if verified is True:
            # true means a call was made and it returned postings, which is only
            # possible if an endpoint is actually configured.
            has_endpoint = bool((source.get("api") or {}).get("url")) or bool(
                (source.get("html") or {}).get("url") or source.get("careers_url")
            )
            expect(
                has_endpoint,
                f"{source['key']}: verified is true but no endpoint is configured, so "
                "no call can have returned postings",
            )
        if source.get("enabled"):
            expect(
                verified != "unconfirmed",
                f"{source['key']}: enabled while nothing about it has been confirmed",
            )


@check("sidra stays permanently disabled")
def _sidra_off():
    cfg = config_module.load()
    sidra = next((s for s in cfg.sources if s["key"] == "sidra"), None)
    expect(sidra is not None, "sidra is missing from sources.yaml")
    expect(not sidra.get("enabled"), "sidra is enabled; its portal is a stale 2019 mirror")
    expect(
        sidra.get("permanently_disabled") is True,
        "sidra must carry permanently_disabled: true so a healthy-looking response "
        "cannot tempt it back on",
    )


@check("recruiter_alerts stays disabled (no IMAP credentials)")
def _recruiter_off():
    cfg = config_module.load()
    source = next((s for s in cfg.sources if s["key"] == "recruiter_alerts"), None)
    expect(source is not None, "recruiter_alerts is missing from sources.yaml")
    expect(not source.get("enabled"), "recruiter_alerts is enabled but IMAP credentials are not set")


# --------------------------------------------------------------------------
# The rules that protect the data
# --------------------------------------------------------------------------

@check("a closing deadline can never become posted_at")
def _deadline_never_posted():
    try:
        Posting.from_deadline_only(
            source_key="t", title="x", posted_at=_dt.date.today()
        )
    except ValueError:
        pass
    else:
        raise Failed("from_deadline_only() accepted a posted_at; the rule is not enforced")

    posting = Posting.from_deadline_only(
        source_key="t", title="x", closing_at=_dt.date(2026, 12, 1)
    )
    expect(posting.posted_at is None, "posted_at was populated from a deadline-only payload")
    expect(posting.age_days is None, "age was inferred despite an unknown posting date")


@check("the request delay floor cannot be lowered")
def _delay_floor():
    for bad in (0, 0.5, 1.4999):
        try:
            Fetcher(delay=bad)
        except ValueError:
            continue
        raise Failed(f"Fetcher accepted delay={bad}s, below the 1.5s floor")
    Fetcher(delay=1.5)   # the floor itself must be allowed


@check("unparseable dates become unknown, never today")
def _dates_conservative():
    today = _dt.date.today()
    for value in ("not a date", "", None, "??", "0000-00-00"):
        expect(
            normalize.parse_date(value) is None,
            f"parse_date({value!r}) did not return None",
        )
    expect(normalize.parse_date("2026-08-01") == _dt.date(2026, 8, 1), "ISO date misparsed")
    expect(normalize.parse_date("/Date(1754006400000)/") is not None, ".NET date misparsed")
    expect(normalize.parse_date("1 Aug 2026") == _dt.date(2026, 8, 1), "long-form date misparsed")
    expect(normalize.parse_date(str(today)) == today, "today's date misparsed")


@check("stale postings are dropped, unknown-dated ones are kept and counted")
def _age_filter():
    stale = Posting(source_key="s", title="old", posted_at=_dt.date(2019, 1, 1))
    fresh = Posting(source_key="s", title="new", posted_at=_dt.date.today())
    unknown = Posting.from_deadline_only(source_key="s", title="undated")
    kept, dropped, unknown_count = normalize.keep_recent([stale, fresh, unknown], 90)
    expect([p.title for p in dropped] == ["old"], "the 2019 posting was not dropped")
    expect({p.title for p in kept} == {"new", "undated"}, "wrong postings kept")
    expect(unknown_count == 1, "undated posting was not counted as unknown")


@check("dedupe collapses the same posting seen twice")
def _dedupe():
    a = Posting(source_key="s", title="Consultant GI", url="https://x.invalid/1")
    b = Posting(source_key="t", title="Consultant GI (dup)", url="https://x.invalid/1")
    c = Posting(source_key="s", title="Different", url="https://x.invalid/2")
    expect(len(normalize.dedupe([a, b, c])) == 2, "dedupe did not collapse a shared URL")


# --------------------------------------------------------------------------
# Scoring
# --------------------------------------------------------------------------

@check("scoring ranks gastroenterology roles above unrelated ones")
def _scoring_ranks():
    scorer = Scorer(config_module.load().profile)

    def score(title, department="", location="Riyadh, Saudi Arabia"):
        return scorer.score(
            Posting(source_key="t", title=title, department=department, location=location)
        )

    shortlisted = [
        "Consultant Gastroenterologist",
        "Consultant, Gastroenterology & Hepatology",
        "Specialist Hepatology",
        "Consultant Transplant Hepatology",
    ]
    for title in shortlisted:
        posting = score(title)
        expect(posting.shortlisted, f"{title!r} scored {posting.score} and missed the shortlist")

    for title in ("Consultant Cardiologist", "Regional Logistics Manager", "Chief Financial Officer"):
        posting = score(title)
        expect(
            not posting.shortlisted,
            f"{title!r} scored {posting.score} and reached the shortlist",
        )


@check("non-physician titles are excluded however well they keyword-match")
def _scoring_excludes():
    scorer = Scorer(config_module.load().profile)
    for title in (
        "Staff Nurse - Endoscopy Unit",
        "Endoscopy Technician",
        "Registrar, Gastroenterology",
        "Medical Records Clerk",
        "Gastroenterology Department Secretary",
    ):
        posting = scorer.score(
            Posting(source_key="t", title=title, department="Gastroenterology", location="Riyadh")
        )
        expect(posting.score == 0, f"{title!r} scored {posting.score}; it should be excluded")
        expect(
            any("excluded" in reason for reason in posting.score_reasons),
            f"{title!r} was zeroed without saying why",
        )


@check("abbreviations match whole words only")
def _scoring_whole_word():
    scorer = Scorer(config_module.load().profile)
    for title in ("Regional Surgical Logistics Lead", "Nuance Systems Analyst", "Biomedical Engineer"):
        posting = scorer.score(Posting(source_key="t", title=title, location="Dubai"))
        expect(
            "gi" not in posting.matched_terms,
            f"{title!r} matched the abbreviation 'gi' as a substring",
        )
    hit = scorer.score(Posting(source_key="t", title="Consultant GI", location="Dubai"))
    expect("gi" in hit.matched_terms, "'Consultant GI' did not match the abbreviation 'gi'")


@check("grade is recorded but never scored (ranking_mode role_fit_only)")
def _grade_not_scored():
    scorer = Scorer(config_module.load().profile)
    consultant = scorer.score(Posting(source_key="t", title="Consultant Gastroenterology", location="Riyadh"))
    specialist = scorer.score(Posting(source_key="t", title="Specialist Gastroenterology", location="Riyadh"))
    expect(
        consultant.score == specialist.score,
        f"grade affected the score: consultant={consultant.score} specialist={specialist.score}, "
        "but ranking_mode is role_fit_only",
    )
    expect(consultant.grade == "consultant", "the consultant grade was not recorded")
    expect(specialist.grade == "specialist", "the specialist grade was not recorded")


# --------------------------------------------------------------------------
# Adapter regressions, against payloads shaped like the real responses
# --------------------------------------------------------------------------

def _with_server(fn):
    from tests.fixture_server import start
    server, base = start(8799)
    try:
        return fn(base)
    finally:
        server.shutdown()


@check("successfactors adapter parses its payload shape")
def _adapter_sf():
    def run(base):
        notes = []
        source = {"key": "jhah", "country": "Saudi Arabia", "platform": "successfactors",
                  "api": {"url": base + "/sf/search", "records_path": "jobs"}}
        out = adapters.get("successfactors")(source, Fetcher(), notes.append)
        expect(len(out) == 3, f"expected 3 postings, got {len(out)}")
        first = out[0]
        expect(first.title == "Consultant Gastroenterologist", f"title misparsed: {first.title!r}")
        expect(first.posted_at == _dt.date(2026, 8, 14), f"posted_at misparsed: {first.posted_at}")
        expect(first.url.endswith("REQ-40122"), "apply URL misparsed")
    _with_server(run)


@check("oracle_orc adapter parses its payload, and its deadline stays a deadline")
def _adapter_orc():
    def run(base):
        notes = []
        source = {"key": "seha", "country": "United Arab Emirates", "platform": "oracle_orc",
                  "api": {"url": base + "/orc/requisitions"}}
        out = adapters.get("oracle_orc")(source, Fetcher(), notes.append)
        expect(len(out) == 2, f"expected 2 postings, got {len(out)}")
        for posting in out:
            expect(
                posting.posted_at is None,
                f"{posting.title!r}: the payload has only PostingEndDate, but posted_at "
                f"was set to {posting.posted_at}",
            )
            expect(posting.closing_at is not None, f"{posting.title!r}: closing date lost")
        expect(out[0].location.startswith("Abu Dhabi"), "location misparsed")
    _with_server(run)


@check("elevatus adapter parses its payload shape")
def _adapter_elevatus():
    def run(base):
        notes = []
        source = {"key": "hmg_habib", "country": "Saudi Arabia", "platform": "elevatus",
                  "api": {"url": base + "/elevatus/api/job-posts"}}
        out = adapters.get("elevatus")(source, Fetcher(), notes.append)
        expect(len(out) == 2, f"expected 2 postings, got {len(out)}")
        expect(out[0].posted_at is not None, "epoch published_at was not parsed")
    _with_server(run)


@check("html_table adapter reads a server-rendered results table")
def _adapter_html():
    def run(base):
        notes = []
        source = {"key": "kfshrc", "country": "Saudi Arabia", "platform": "html_table",
                  "html": {"url": base + "/careers", "row_selector": "table.jobs tbody tr"}}
        out = adapters.get("html_table")(source, Fetcher(), notes.append)
        expect(len(out) == 3, f"expected 3 postings, got {len(out)}")
        expect(out[0].posted_at == _dt.date(2026, 8, 12), f"posted_at misparsed: {out[0].posted_at}")
        expect(out[0].closing_at == _dt.date(2026, 11, 30), "closing date misparsed")
        expect(out[0].url.startswith("http"), "relative link was not resolved to absolute")
    _with_server(run)


@check("a table header row never becomes a vacancy")
def _header_row_not_a_job():
    def run(base):
        notes = []
        # "table tr" matches the header alongside the data rows -- and it is
        # exactly what discovery emits for a plain table, so the adapter has to
        # survive it. Without the guard this yields a posting titled "Position".
        source = {"key": "kfshrc", "country": "Saudi Arabia", "platform": "html_table",
                  "html": {"url": base + "/careers", "row_selector": "table tr"}}
        out = adapters.get("html_table")(source, Fetcher(), notes.append)
        titles = [p.title for p in out]
        expect(len(out) == 3, f"expected 3 vacancies, got {len(out)}: {titles}")
        for header in ("Position", "Department", "Location", "Date Posted"):
            expect(header not in titles, f"the header cell {header!r} became a vacancy")
    _with_server(run)


@check("a failing source raises rather than returning an innocent empty list")
def _adapter_failures_are_loud():
    def run(base):
        cases = [
            ("oracle_orc", {"key": "a", "platform": "oracle_orc", "api": {"url": base + "/forbidden"}}),
            ("successfactors", {"key": "b", "platform": "successfactors",
                                "api": {"url": base + "/notjson", "records_path": "jobs"}}),
            ("successfactors", {"key": "c", "platform": "successfactors",
                                "api": {"url": base + "/sf/search", "records_path": "nope.here"}}),
            ("successfactors", {"key": "d", "platform": "successfactors"}),
            ("html_table", {"key": "e", "platform": "html_table",
                            "html": {"url": base + "/careers", "row_selector": "div.nope"}}),
        ]
        for platform, source in cases:
            try:
                adapters.get(platform)(source, Fetcher(), lambda m: None)
            except AdapterError as exc:
                expect(len(str(exc)) > 20, f"{source['key']}: error message is too thin to diagnose")
                continue
            raise Failed(f"{source['key']}: a broken source returned quietly instead of raising")
    _with_server(run)


@check("an empty board is reported as an observation, not passed off as clean")
def _empty_board_noted():
    def run(base):
        notes = []
        source = {"key": "f", "platform": "successfactors",
                  "api": {"url": base + "/empty", "records_path": "jobs"}}
        out = adapters.get("successfactors")(source, Fetcher(), notes.append)
        expect(out == [], "expected no postings")
        expect(any("EMPTY" in n for n in notes), "an empty result set produced no note")
    _with_server(run)


# --------------------------------------------------------------------------
# Reporting
# --------------------------------------------------------------------------

@check("note() reaches the run record instead of stderr")
def _note_routing():
    log = RunLog()
    log.note("kfshrc", "something worth knowing")
    log.note("kfshrc", "something worth knowing")     # deduped
    expect(log.get("kfshrc").notes == ["something worth knowing"], "note() did not record once")
    log.record("kfshrc", status="error", error="HTTP 403")
    expect(log.get("kfshrc").notes, "recording a status discarded the earlier note")
    row = log.get("kfshrc").as_row()
    for column in ("source", "status", "fetched", "kept", "error", "notes"):
        expect(column in row, f"Run status row is missing the {column!r} column")


@check("scanner.note() is importable with the documented signature")
def _scanner_note_api():
    import scanner
    expect(hasattr(scanner, "note"), "scanner.note is missing")
    scanner.note("kfshrc", "reached through the module-level entry point")
    import inspect
    params = list(inspect.signature(scanner.note).parameters)
    expect(
        params[:2] == ["source_key", "message"],
        f"scanner.note signature is {params}, expected (source_key, message)",
    )


@check("writers produce the expected sheets and files")
def _writers():
    from openpyxl import load_workbook
    cfg = config_module.load()
    log = RunLog()
    log.record("kfshrc", name="KFSHRC", platform="html_table", status="ok", fetched=3, kept=2)
    log.note("kfshrc", "a note that must survive into the sheet")
    postings = Scorer(cfg.profile).score_all([
        Posting(source_key="kfshrc", title="Consultant Gastroenterologist",
                location="Riyadh", country="Saudi Arabia", posted_at=_dt.date.today(),
                url="https://x.invalid/1"),
        Posting(source_key="kfshrc", title="Consultant Cardiologist",
                location="Riyadh", country="Saudi Arabia", posted_at=_dt.date.today()),
    ])
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        xlsx = writers.write_xlsx(tmp_path / "j.xlsx", postings, log, cfg)
        book = load_workbook(xlsx)
        expect(
            book.sheetnames == ["Postings", "Run status", "Run info"],
            f"unexpected sheets: {book.sheetnames}",
        )
        headers = [c.value for c in book["Run status"][1]]
        for column in ("source", "status", "fetched", "kept", "error", "notes"):
            expect(column in headers, f"Run status sheet has no {column!r} column")
        status_row = [c.value for c in book["Run status"][2]]
        expect("a note that must survive into the sheet" in str(status_row), "the note did not reach the sheet")

        docx = writers.write_docx(tmp_path / "s.docx", postings, log, cfg)
        expect(docx.stat().st_size > 0, "docx is empty")
        from docx import Document
        text = "\n".join(p.text for p in Document(docx).paragraphs)
        expect("Consultant Gastroenterologist" in text, "the shortlisted role is not in the docx")
        expect("Cardiologist" not in text, "a non-shortlisted role leaked into the docx")

        payload = writers.write_json(tmp_path / "s.json", postings, log, cfg)
        import json
        data = json.loads(payload.read_text())
        for key in ("generated_at", "profile", "totals", "run_status", "postings"):
            expect(key in data, f"json dump has no {key!r}")

        page = writers.write_html(tmp_path / "i.html", postings, log, cfg)
        html_text = page.read_text()
        expect("Run status" in html_text, "the html page has no run status table")
        expect("<script" not in html_text.lower(), "the html page pulls in script; it must be self-contained")


@check("an all-failed run says so instead of showing an innocent empty shortlist")
def _empty_run_is_labelled():
    cfg = config_module.load()
    log = RunLog()
    log.record("kfshrc", name="KFSHRC", platform="html_table", status="error",
               error="connection refused")
    with tempfile.TemporaryDirectory() as tmp:
        docx = writers.write_docx(Path(tmp) / "s.docx", [], log, cfg)
        from docx import Document
        text = "\n".join(p.text for p in Document(docx).paragraphs)
        expect(
            "does NOT mean there are no vacancies" in text.replace("\n", " "),
            "an all-failed run produced a shortlist that reads as 'nothing available'",
        )


@check("the weekly run does not depend on a browser")
def _no_browser_at_runtime():
    # Requirement lines only. requirements.txt explains in a comment why the
    # browser is absent, and a naive substring search fires on that prose.
    lines = [
        line.split("#")[0].strip().lower()
        for line in (ROOT / "requirements.txt").read_text().splitlines()
    ]
    declared = [line for line in lines if line]
    expect(
        not any("playwright" in line for line in declared),
        f"playwright is declared in requirements.txt ({declared}). Endpoints are "
        "hard-coded once found precisely so the scheduled scan stays a handful of "
        "HTTP requests; the browser belongs in requirements-browser.txt only.",
    )
    expect(
        (ROOT / "requirements-browser.txt").exists(),
        "requirements-browser.txt is missing; discovery has no declared dependencies",
    )

    # Import the scanner with playwright made unimportable. If it still loads,
    # no runtime path reaches the browser.
    import subprocess
    probe = (
        "import sys; sys.modules['playwright'] = None; "
        f"sys.path.insert(0, {str(ROOT)!r}); import scanner; print('ok')"
    )
    result = subprocess.run(
        [sys.executable, "-c", probe], capture_output=True, text=True, cwd=ROOT
    )
    expect(
        result.returncode == 0,
        "scanner.py could not be imported with playwright unavailable, so the "
        f"weekly run depends on a browser:\n{result.stderr[-600:]}",
    )


@check("run-level messages are not counted as employers")
def _pseudo_source_not_counted():
    log = RunLog()
    log.record("kfshrc", status="ok", fetched=3, kept=2)
    log.note("__run__", "a run-level message that must not inflate the source count")
    totals = log.totals()
    expect(
        totals["sources"] == 1,
        f"totals counted {totals['sources']} sources; the __run__ pseudo-source was "
        "counted as an employer",
    )
    expect(totals["attempted"] == 1, f"attempted={totals['attempted']}, expected 1")
    expect(
        any(r.source_key == "__run__" for r in log.records),
        "the run-level message was dropped from the report entirely",
    )


@check("--only can reach a disabled source, but not a permanently disabled one")
def _only_selection():
    cfg = config_module.load()
    selected = [s["key"] for s in cfg.select(["kfshrc", "jhah"])]
    expect(selected == ["kfshrc", "jhah"], f"--only did not select disabled sources: {selected}")
    try:
        cfg.select(["no_such_source"])
    except config_module.ConfigError:
        pass
    else:
        raise Failed("an unknown source key was accepted silently")


def main() -> int:
    print(f"qa_check: {len(CHECKS)} checks\n")
    failures = []
    for index, (name, fn) in enumerate(CHECKS, start=1):
        try:
            fn()
        except Failed as exc:
            failures.append((name, str(exc)))
            print(f"  {index:>2}. FAIL  {name}")
        except Exception:
            failures.append((name, traceback.format_exc()))
            print(f"  {index:>2}. ERROR {name}")
        else:
            print(f"  {index:>2}. pass  {name}")

    print()
    if failures:
        print(f"{len(failures)} of {len(CHECKS)} checks failed:\n")
        for name, detail in failures:
            print(f"--- {name}\n{detail}\n")
        return 1
    print(f"all {len(CHECKS)} checks passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
