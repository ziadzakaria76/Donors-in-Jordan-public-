"""The shortlist document: the roles worth acting on, and why.

Each entry carries the reasons the scorer recorded. A shortlist that cannot
explain itself is a shortlist that gets re-checked by hand, which defeats the
point of running the scan at all.
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x38, 0x64)
MUTED = RGBColor(0x59, 0x59, 0x59)


def _meta_line(document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def write_docx(path, postings, run_log, config) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    shortlisted = sorted(
        [p for p in postings if p.shortlisted], key=lambda p: (-p.score, p.source_key)
    )
    totals = run_log.totals()

    document = Document()
    heading = document.add_heading("Shortlist", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _meta_line(
        document,
        f"{config.profile.get('target', 'target role')} — generated "
        f"{_dt.date.today().isoformat()}",
    )
    _meta_line(
        document,
        f"Scored {len(postings)} posting(s) from {totals['sources']} source(s); "
        f"{totals['ok']} returned data, {totals['error'] + totals['blocked']} failed. "
        f"Shortlist threshold: score ≥ {config.shortlist_min_score}.",
    )

    if totals["ok"] == 0:
        document.add_paragraph()
        warning = document.add_paragraph()
        run = warning.add_run(
            "No source returned any postings in this run. An empty shortlist here "
            "means the scan could not read any employer — it does NOT mean there "
            "are no vacancies. See the Run status sheet in the accompanying "
            "workbook for the failure against each source."
        )
        run.bold = True

    document.add_paragraph()

    if not shortlisted:
        document.add_paragraph(
            "Nothing met the shortlist threshold in this run."
        )
    for posting in shortlisted:
        title = document.add_heading(posting.title, level=2)
        for run in title.runs:
            run.font.color.rgb = ACCENT

        facts = []
        if posting.location:
            facts.append(posting.location)
        if posting.department:
            facts.append(posting.department)
        facts.append(f"{posting.source_key} · score {posting.score}")
        if posting.posted_at:
            facts.append(f"posted {posting.posted_at.isoformat()}")
        else:
            facts.append("posting date not published")
        if posting.closing_at:
            facts.append(f"closes {posting.closing_at.isoformat()}")
        _meta_line(document, "  ·  ".join(facts))

        for reason in posting.score_reasons:
            bullet = document.add_paragraph(reason, style="List Bullet")
            for run in bullet.runs:
                run.font.size = Pt(9)

        if posting.url:
            link = document.add_paragraph()
            run = link.add_run(posting.url)
            run.font.size = Pt(9)
            run.font.color.rgb = ACCENT

    document.save(target)
    return target
