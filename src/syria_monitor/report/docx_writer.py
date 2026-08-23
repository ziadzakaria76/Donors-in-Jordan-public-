"""Word report -- the bid-review pack.

Arabic titles get TRUE right-to-left paragraph direction: <w:bidi/> on the
paragraph properties and <w:rtl/> on every run, plus a complex-script font.
Right alignment alone is not RTL -- it moves the text to the right margin while
leaving the reading order, punctuation placement and last-line behaviour wrong.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..models import LINK_TYPES
from .common import (LINK_LABELS, SCREENING_DISCLAIMER, badges, fmt_date, fmt_value,
                     split_live_pipeline)

ARABIC_FONT = "Arial"


def _is_arabic(text: str) -> bool:
    return any("؀" <= ch <= "ۿ" for ch in text or "")


def _apply_rtl(paragraph) -> None:
    """Real RTL: w:bidi on the paragraph, w:rtl + w:cs on each run."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        for tag in ("w:rtl", "w:cs"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), "1")
            rPr.append(el)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:cs"), ARABIC_FONT)


def _para(doc, text, size=10.5, bold=False, color=None, space_after=4, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text or "")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if _is_arabic(text):
        _apply_rtl(paragraph)
    return paragraph


def write_docx(result, path: Path, top_n: int = 10) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading("Syria Tender Intelligence", level=0)
    _para(doc, result.subject(), size=11, bold=True, color="1B4F72")
    _para(doc, f"Run started {result.started} | in scope: {len(result.tenders)} "
               f"| new: {len(result.new_tenders)} | excluded but logged: {len(result.excluded)}",
          size=9, color="6B7280")

    # 1 --- executive brief
    doc.add_heading(f"1. Top {min(top_n, len(result.tenders))} this run", level=1)
    if not result.tenders:
        _para(doc, "No tenders in scope this run. See run diagnostics below for portal health "
                   "-- a quiet day and a broken scraper look identical without it.", bold=True)
    for rank, tender in enumerate(result.tenders[:top_n], start=1):
        _para(doc, f"{rank}. [{tender.score:.0f}] {tender.title}", size=11, bold=True)
        _para(doc, f"     {tender.portal} | closes {fmt_date(tender.closing_date)} | "
                   f"{fmt_value(tender)} | {LINK_LABELS.get(tender.syria_link_type, '')}", size=9)
        marks = badges(tender)
        if marks:
            _para(doc, "     " + " · ".join(marks), size=9, color="7D3C00")

    # 2 --- full list
    live, pipeline = split_live_pipeline(result.tenders)
    doc.add_heading("2. Full list", level=1)
    _section_table(doc, f"LIVE - biddable ({len(live)})", live)
    _section_table(doc, f"PIPELINE - GPN / advance notices, not yet biddable ({len(pipeline)})",
                   pipeline)

    # 3 --- detail
    doc.add_heading("3. Detail", level=1)
    for tender in result.tenders:
        _para(doc, tender.title, size=11, bold=True)
        _para(doc, f"{tender.portal} | {tender.notice_type or 'type not stated'} | "
                   f"posted {fmt_date(tender.posted_date)} | closes {fmt_date(tender.closing_date)}",
              size=9, color="6B7280")
        if tender.description:
            _para(doc, tender.description[:1200], size=9.5)
        if tender.eligibility:
            _para(doc, f"Eligibility: {tender.eligibility}", size=9)
        if tender.contact:
            _para(doc, f"Contact: {tender.contact}", size=9)
        if tender.screening:
            _para(doc, "Sanctions flag: " + "; ".join(
                f"{h['party']} ~ {h['matched_name']} ({h['list']}, list fetched {h['list_fetched']})"
                for h in tender.screening), size=9, bold=True, color="B02A2A")
            _para(doc, SCREENING_DISCLAIMER, size=8.5, color="B02A2A")
        _para(doc, f"Value: {fmt_value(tender)}", size=9)
        _para(doc, tender.url or "no link published (id did not match the notice-id pattern)",
              size=9, color="1B4F72", space_after=10)

    # 4 --- diagnostics
    doc.add_heading("4. Run diagnostics", level=1)
    _para(doc, "Portal health", bold=True)
    for portal in result.portals:
        colour = "B02A2A" if not portal.available else ("6B7280" if portal.skipped_reason else "1F2328")
        _para(doc, "  " + portal.status_line, size=9, color=colour)

    _para(doc, "Classification split (all categories, including those out of scope)",
          bold=True, space_after=2)
    for key in LINK_TYPES:
        _para(doc, f"  {LINK_LABELS.get(key, key)}: {result.counts.get(key, 0)}", size=9)
    _para(doc, f"  duplicates collapsed across portals: {result.duplicates_collapsed} | "
               f"expired dropped: {result.expired_dropped}", size=9)

    _para(doc, "Sanctions screening", bold=True, space_after=2)
    if result.screening_error:
        _para(doc, f"  ERROR: {result.screening_error}", size=9, color="B02A2A")
    for entry in result.screening_status:
        _para(doc, f"  {entry['list']}: fetched {entry['fetched']}, {entry['names']} names"
                   + (f" -- {entry['error']}" if entry.get("error") else ""), size=9)
    _para(doc, SCREENING_DISCLAIMER, size=9, bold=True, color="B02A2A")

    doc.save(path)
    return path


def _section_table(doc, heading: str, tenders: list) -> None:
    _para(doc, heading, size=10.5, bold=True, space_after=2)
    if not tenders:
        _para(doc, "  none", size=9, color="6B7280")
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for idx, label in enumerate(("", "Title", "Portal", "Closes", "Value")):
        table.rows[0].cells[idx].text = label
    for tender in tenders:
        cells = table.add_row().cells
        cells[0].text = "NEW" if tender.is_new else ""
        cells[1].text = tender.title[:160]
        if _is_arabic(tender.title):
            for paragraph in cells[1].paragraphs:
                _apply_rtl(paragraph)
        cells[2].text = tender.portal
        cells[3].text = fmt_date(tender.closing_date)
        cells[4].text = fmt_value(tender)
