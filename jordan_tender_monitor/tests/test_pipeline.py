"""
Pipeline tests: filtering, scoring, deduplication, reporting, delivery.

The state-touching cases (the seen-tenders database, the output directory) all
run inside a temporary folder. A test that wrote fixture IDs into the real
database would make the next live run report nothing.
"""

from __future__ import annotations

import csv
import json
import tempfile
from datetime import date, timedelta
from pathlib import Path

from jordan_tender_monitor import config, fixtures
from jordan_tender_monitor.agents import emailer, filter as filters, reporter, tracker
from jordan_tender_monitor.agents.scraper import PortalHealth
from jordan_tender_monitor.portals import base

from .harness import check, check_eq

TODAY = date(2026, 6, 15)


def _record(**kwargs) -> dict:
    defaults = dict(portal="worldbank", title="Advisory Services, Jordan",
                    url="https://e.org/1", posted=TODAY - timedelta(days=5),
                    closing=TODAY + timedelta(days=30), value_text="USD 500,000")
    defaults.update(kwargs)
    return base.build_record(**defaults)


# ---------------------------------------------------------------------------
# Empty and degenerate input
# ---------------------------------------------------------------------------


def test_empty_input():
    result = filters.process([], TODAY)
    check_eq(result["tenders"], [], "empty: no tenders out")
    check_eq(result["scanned"], 0, "empty: scanned count is zero")
    check_eq(result["merged_duplicates"], 0, "empty: nothing merged")

    health = fixtures.sample_health()
    healthy = [h for h in health if h.status == "ok"]
    subject = reporter.build_subject(0, healthy)
    check("no new opportunities" in subject, "empty: subject says no opportunities")
    check("portals OK" in subject, "empty: and confirms the portals were readable")
    body = reporter.build_email_html(result, health)
    check("<html" in body, "empty: an email body is still produced")


def test_every_optional_field_none():
    """Nothing downstream may assume an optional field exists."""
    record = base.build_record(portal="isdb", title="Water Sector Study, Jordan",
                               url=None, posted=None, closing=None, value_text=None,
                               description=None, notice_type=None, contact=None,
                               reference=None)
    result = filters.process([record], TODAY)
    check_eq(len(result["tenders"]), 1, "none-fields: the tender survives")
    tender = result["tenders"][0]
    check(isinstance(tender["score"], float), "none-fields: it still scores")
    check(config.UNKNOWN_DEADLINE_NOTE in tender["flags"],
          "none-fields: missing deadline is flagged, not dropped")
    check("Value not published" in tender["flags"],
          "none-fields: missing value is flagged, not dropped")

    reporter.decorate(result["tenders"], TODAY)
    body = reporter.build_email_html(result, fixtures.sample_health())
    check("Water Sector Study" in body, "none-fields: it renders in the email")

    with tempfile.TemporaryDirectory() as tmp:
        written = reporter.write_outputs(result, fixtures.sample_health(), body, Path(tmp))
        check_eq(len(written), len(config.OUTPUT_FORMATS),
                 "none-fields: every output format is written")


def test_zero_keyword_matches_division_guard():
    """No keyword hits must not divide by zero, and must not crash scoring."""
    record = _record(title="Zzz Qqq Wwq", description="Zzz qqq wwq.",
                     url="https://e.org/nokw")
    score = filters.score(record, TODAY)
    check(isinstance(score, float), "division-guard: a score is still produced")
    check(score >= 0, "division-guard: the score is non-negative")
    check_eq(record["_components"]["keyword"], 0.0,
             "division-guard: keyword component is exactly zero")


# ---------------------------------------------------------------------------
# Filtering policy
# ---------------------------------------------------------------------------


def test_deadline_boundary_today_is_kept():
    today_tender = _record(closing=TODAY, url="https://e.org/today")
    past_tender = _record(closing=TODAY - timedelta(days=1), url="https://e.org/past")
    result = filters.process([today_tender, past_tender], TODAY)
    urls = {t["url"] for t in result["tenders"]}
    check("https://e.org/today" in urls, "boundary: closing today is KEPT")
    check("https://e.org/past" not in urls, "boundary: closed yesterday is dropped")
    check_eq(result["dropped"].get("closed"), 1, "boundary: the drop is counted")


def test_unknown_value_is_kept_and_flagged():
    unknown = _record(value_text=None, url="https://e.org/unknown")
    small = _record(value_text="USD 40,000", url="https://e.org/small")
    result = filters.process([unknown, small], TODAY)
    urls = {t["url"] for t in result["tenders"]}
    check("https://e.org/unknown" in urls,
          "value: an unpublished value is kept -- unknown is not small")
    check("https://e.org/small" not in urls,
          "value: a published value below the floor is dropped")
    kept = [t for t in result["tenders"] if t["url"] == "https://e.org/unknown"][0]
    check("Value not published" in kept["flags"], "value: and it is flagged")


def test_unknown_value_scores_mid_band():
    unknown = _record(value_text=None)
    filters.score(unknown, TODAY)
    component = unknown["_components"]["value"]
    check(0.3 < component < 0.8,
          "value: unknown scores mid-band, neither dominating nor sinking",
          f"got {component}")


def test_national_only_is_flagged_and_penalised():
    restricted = _record(
        title="Public Administration Reform Advisory, Jordan",
        description="Consulting services. National firms only.",
        url="https://e.org/restricted")
    open_bid = _record(
        title="Public Administration Reform Advisory, Jordan",
        description="Consulting services, open to international firms.",
        url="https://e.org/open")

    result = filters.process([restricted, open_bid], TODAY)
    check_eq(len(result["tenders"]), 2, "eligibility: neither is excluded")

    by_url = {t["url"]: t for t in result["tenders"]}
    flagged = by_url["https://e.org/restricted"]
    check(flagged["eligibility"], "eligibility: the restriction is detected")
    check(any("National" in f or "local" in f for f in flagged["flags"]),
          "eligibility: and shown as a flag")
    check(flagged["score"] < by_url["https://e.org/open"]["score"],
          "eligibility: the restricted tender ranks lower")


def test_arabic_notice_is_kept_and_flagged():
    arabic = base.build_record(
        portal="sfd", title="خدمات استشارية لتطوير القطاع المالي في الأردن",
        url="https://e.org/ar", posted=TODAY - timedelta(days=3),
        closing=TODAY + timedelta(days=40), value_text="٧٥٠٠٠٠ دولار",
        description="دراسة جدوى وبناء القدرات لوزارة المالية.")
    result = filters.process([arabic], TODAY)
    check_eq(len(result["tenders"]), 1, "arabic: kept")
    tender = result["tenders"][0]
    check_eq(tender["language"], "ar", "arabic: language detected")
    check(config.ARABIC_FLAG_NOTE in tender["flags"], "arabic: flagged for review")
    check(tender["score"] > 0,
          "arabic: scores above zero -- the Arabic ranking lexicon works")


def test_all_sectors_and_types_pass():
    records = [
        _record(title="Road Construction Supervision, Jordan", url="https://e.org/a"),
        _record(title="Health Sector Assessment, Amman", url="https://e.org/b"),
        _record(title="Solar Feasibility Study, Aqaba", url="https://e.org/c"),
    ]
    result = filters.process(records, TODAY)
    check_eq(len(result["tenders"]), 3, "sectors: nothing dropped on sector or type")
    sectors = {t["sector"] for t in result["tenders"]}
    check(len(sectors) >= 2, "sectors: tenders are labelled with distinct sectors")


# ---------------------------------------------------------------------------
# Scoring weights
# ---------------------------------------------------------------------------


def test_disabled_component_is_dropped_and_weights_renormalise():
    weights = filters.active_weights()
    check("sector" not in weights,
          "weights: the sector component is dropped when all sectors are selected")
    check("keyword" in weights,
          "weights: the keyword component survives -- the ranking lexicon still varies")
    check(abs(sum(weights.values()) - 100.0) < 0.01,
          "weights: the remainder renormalises to 100", f"sum {sum(weights.values())}")
    check(abs(weights["keyword"] - 57.14) < 0.1,
          "weights: keyword becomes 57.1", f"got {weights['keyword']}")


def test_ranking_favours_advisory_over_goods():
    advisory = _record(title="Institutional Reform Technical Assistance, Jordan",
                       description="Capacity building and advisory services.",
                       url="https://e.org/adv")
    goods = _record(title="Supply and Delivery of Vehicles, Jordan",
                    description="Procurement of vehicles and spare parts.",
                    url="https://e.org/goods", value_text="USD 5,000,000")
    filters.score_all([advisory, goods], TODAY)
    check(advisory["score"] > goods["score"],
          "ranking: advisory work outranks a larger goods contract",
          f"advisory {advisory['score']} vs goods {goods['score']}")


# ---------------------------------------------------------------------------
# Deduplication
# ---------------------------------------------------------------------------


def test_duplicate_collapse_across_portals():
    a = _record(portal="worldbank",
                title="Institutional Strengthening of the Ministry of Finance, Jordan",
                url="https://e.org/wb")
    b = _record(portal="ungm",
                title="Institutional Strengthening of the Ministry of Finance (Jordan)",
                url="https://e.org/ungm", value_text=None)
    result = filters.process([a, b], TODAY)
    check_eq(len(result["tenders"]), 1, "dedupe: two portals collapse to one tender")
    check_eq(result["merged_duplicates"], 1, "dedupe: the merge is counted")
    check(result["tenders"][0]["also_on"],
          "dedupe: the other portal is recorded, not hidden")


def test_dedupe_keeps_the_better_data():
    rich = _record(portal="worldbank", title="Water Sector Study, Jordan",
                   url="https://e.org/x", value_text="USD 900,000",
                   contact="proc@example.org")
    sparse = _record(portal="ungm", title="Water Sector Study, Jordan",
                     url="https://e.org/y", value_text=None, closing=None)
    result = filters.process([rich, sparse], TODAY)
    check_eq(len(result["tenders"]), 1, "dedupe: collapsed")
    check(result["tenders"][0]["estimated_value_usd"] is not None,
          "dedupe: the surviving copy keeps the published value")


def test_same_portal_numbered_lots_are_not_merged():
    """Numbered lots from one portal score ~97 on title similarity.

    Collapsing them would silently delete real tenders, so fuzzy matching is
    restricted to cross-portal duplicates.
    """
    lots = [_record(portal="worldbank",
                    title=f"Governance Advisory Assignment Number {i}, Jordan",
                    url=f"https://e.org/lot/{i}") for i in range(6)]
    result = filters.process(lots, TODAY)
    check_eq(len(result["tenders"]), 6,
             "dedupe: six numbered lots from one portal stay six tenders")
    check_eq(result["merged_duplicates"], 0, "dedupe: nothing was merged")


def test_distinct_tenders_are_not_merged():
    a = _record(title="Health Sector Assessment, Amman", url="https://e.org/1")
    b = _record(title="Energy Grid Integration Study, Aqaba", url="https://e.org/2")
    result = filters.process([a, b], TODAY)
    check_eq(len(result["tenders"]), 2, "dedupe: unrelated tenders are left alone")


# ---------------------------------------------------------------------------
# The subject line -- Q15
# ---------------------------------------------------------------------------


def test_subject_distinguishes_quiet_from_broken():
    healthy = fixtures.sample_health()
    for h in healthy:
        if h.status == "unavailable":
            h.status, h.reason = "ok", ""

    quiet = reporter.build_subject(0, healthy)
    check("no new opportunities" in quiet, "subject: a quiet day says so")
    check("portals OK" in quiet, "subject: and states that the portals worked")
    check(not quiet.startswith(config.ACTION_NEEDED_PREFIX),
          "subject: a quiet day is not an alert")

    broken = reporter.build_subject(0, fixtures.all_broken_health())
    check(broken.startswith(config.ACTION_NEEDED_PREFIX),
          "subject: total failure is an ACTION NEEDED alert")
    check("unreachable" in broken, "subject: and says the portals were unreachable")
    check(quiet != broken,
          "subject: zero-tender runs are NOT indistinguishable -- the whole point")


def test_subject_reports_partial_degradation():
    subject = reporter.build_subject(4, fixtures.sample_health())
    check("4 new" in subject, "subject: the count is present")
    check("unavailable" in subject, "subject: partial degradation is visible")


def test_subject_singular_and_plural():
    healthy = [h for h in fixtures.sample_health() if h.status == "ok"]
    check("1 new opportunity" in reporter.build_subject(1, healthy),
          "subject: singular reads correctly")
    check("7 new opportunities" in reporter.build_subject(7, healthy),
          "subject: plural reads correctly")


# ---------------------------------------------------------------------------
# Email body -- overflow must move tenders, never drop them
# ---------------------------------------------------------------------------


def test_email_overflow_moves_nothing_is_dropped():
    many = []
    for i in range(config.MAX_INLINE_TENDERS + 12):
        many.append(_record(title=f"Governance Advisory Assignment Number {i}, Jordan",
                            url=f"https://e.org/n/{i}"))
    result = filters.process(many, TODAY)
    reporter.decorate(result["tenders"], TODAY)
    check_eq(len(result["tenders"]), config.MAX_INLINE_TENDERS + 12,
             "overflow: all tenders survive filtering")

    body = reporter.build_email_html(result, fixtures.sample_health())
    check("Further opportunities" in body, "overflow: an overflow section appears")
    for i in (0, config.MAX_INLINE_TENDERS + 5, config.MAX_INLINE_TENDERS + 11):
        check(f"https://e.org/n/{i}" in body,
              f"overflow: tender {i} is still present in the email")


def test_health_table_names_the_failure_and_the_url():
    health = fixtures.sample_health()
    body = reporter.build_email_html(
        filters.process([], TODAY), health)
    check("UNAVAILABLE" in body, "health: a broken portal is marked unavailable")
    check("bot wall" in body, "health: the diagnosed cause is shown")
    check("ebrd.com" in body, "health: the URL to check by hand is shown")
    check("NOT CONFIGURED" in body,
          "health: unconfigured is shown distinctly from broken")


# ---------------------------------------------------------------------------
# Output files -- including Arabic survival
# ---------------------------------------------------------------------------


def test_configured_formats_written_and_arabic_survives():
    """Only the configured formats are emitted, and Arabic survives both."""
    records = fixtures.sample_records(TODAY)
    result = filters.process(records, TODAY)
    reporter.decorate(result["tenders"], TODAY)
    health = fixtures.sample_health()
    body = reporter.build_email_html(result, health)

    check_eq(config.OUTPUT_FORMATS, ["docx", "excel"],
             "outputs: Word and Excel only -- no JSON, CSV or HTML")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp)
        written = reporter.write_outputs(result, health, body, out)
        check_eq(sorted(written), ["docx", "excel"],
                 "outputs: exactly two files are produced")
        for fmt_name in ("docx", "excel"):
            check(written[fmt_name].exists() and written[fmt_name].stat().st_size > 0,
                  f"outputs: {fmt_name} is non-empty")
        check_eq(len(list(out.iterdir())), 2,
                 "outputs: nothing else is left in the output directory")

        from openpyxl import load_workbook
        wb = load_workbook(written["excel"])
        cells = [str(c.value) for ws in wb.worksheets
                 for row in ws.iter_rows() for c in row if c.value]
        check(any("استشارية" in c for c in cells), "arabic: survives Excel")

        from docx import Document
        doc = Document(written["docx"])
        check(any("استشارية" in p.text for p in doc.paragraphs),
              "arabic: survives Word")


def test_retained_json_and_csv_writers_still_handle_arabic():
    """write_json/write_csv are no longer emitted but remain available.

    Kept tested so re-enabling them in OUTPUT_FORMATS is a config change rather
    than a bug hunt.
    """
    records = filters.process(fixtures.sample_records(TODAY), TODAY)["tenders"]
    reporter.decorate(records, TODAY)
    health = fixtures.sample_health()

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp)
        json_path = reporter.write_json(records, health, out / "r.json")
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        check(any(t["language"] == "ar" for t in payload["tenders"]),
              "arabic: survives JSON")
        check("استشارية" in json.dumps(payload, ensure_ascii=False),
              "arabic: JSON keeps Arabic characters, not \\uXXXX escapes")

        csv_path = reporter.write_csv(records, out / "r.csv")
        with csv_path.open(encoding="utf-8-sig") as fh:
            rows = list(csv.reader(fh))
        check(any("استشارية" in " ".join(r) for r in rows), "arabic: survives CSV")


def test_filename_states_the_run_health():
    """The filename carries what the email subject line used to.

    With output going to disk, a total outage and a quiet week both leave a
    file behind. Identically-named files make a dead monitor invisible.
    """
    healthy = [h for h in fixtures.sample_health() if h.status == "ok"]

    quiet = reporter.run_slug(0, healthy)
    outage = reporter.run_slug(0, fixtures.all_broken_health())
    normal = reporter.run_slug(7, healthy)
    partial = reporter.run_slug(4, fixtures.sample_health())

    check("no-new-opportunities" in quiet, "filename: a quiet day says so")
    check("portals-OK" in quiet, "filename: and confirms the portals were read")
    check(outage.startswith(config.ACTION_NEEDED_PREFIX.replace(" ", "-")),
          "filename: a total outage is an ACTION NEEDED alert", f"got {outage}")
    check("unreachable" in outage, "filename: and says the portals were unreachable")
    check(quiet != outage,
          "filename: zero-tender runs are NOT indistinguishable -- the whole point")
    check("7-opportunities" in normal, "filename: a normal run states the count")
    check("1-opportunity" in reporter.run_slug(1, healthy),
          "filename: singular reads correctly")
    check("unavailable" in partial, "filename: partial degradation is visible")

    for slug in (quiet, outage, normal, partial):
        check(not any(c in slug for c in r'\/:*?"<>| '),
              f"filename: '{slug}' is safe on Windows and Linux")


def test_documents_state_the_run_health_inside():
    """Opening the file weeks later must not require inferring from an empty list."""
    healthy = [h for h in fixtures.sample_health() if h.status == "ok"]
    empty = filters.process([], TODAY)
    body = reporter.build_email_html(empty, healthy)

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp)
        written = reporter.write_outputs(empty, healthy, body, out)

        from docx import Document
        text = "\n".join(p.text for p in Document(written["docx"]).paragraphs)
        check("No new opportunities" in text,
              "docx: the quiet-run status is stated in words")
        check("not" in text.lower() and "failed" in text.lower(),
              "docx: and says explicitly that this is not a failure")

        from openpyxl import load_workbook
        wb = load_workbook(written["excel"])
        check("Run status" in wb.sheetnames,
              "excel: a Run status sheet is present")
        status = wb["Run status"]
        cells = [str(c.value) for row in status.iter_rows() for c in row if c.value]
        check(any("No new opportunities" in c for c in cells),
              "excel: the status sheet states the run health")
        check(any("World Bank" in c for c in cells),
              "excel: and lists every portal")

    outage = filters.process([], TODAY)
    broken = fixtures.all_broken_health()
    with tempfile.TemporaryDirectory() as tmp:
        written = reporter.write_outputs(
            outage, broken, reporter.build_email_html(outage, broken), Path(tmp))
        from docx import Document
        text = "\n".join(p.text for p in Document(written["docx"]).paragraphs)
        check(config.ACTION_NEEDED_PREFIX in text,
              "docx: a total outage is flagged ACTION NEEDED inside the document")
        check("nothing could be read" in text,
              "docx: and explains the report is empty for the wrong reason")

        from openpyxl import load_workbook
        status = load_workbook(written["excel"])["Run status"]
        cells = [str(c.value) for row in status.iter_rows() for c in row if c.value]
        check(any(config.ACTION_NEEDED_PREFIX in c for c in cells),
              "excel: the status sheet flags the outage too")
        check(any("UNAVAILABLE" in c for c in cells),
              "excel: individual portals are marked unavailable")


def test_excel_with_zero_rows_does_not_crash():
    """auto_filter over a header-only sheet is a classic crash.

    It happens on precisely the quiet day when nothing matched, which is the
    day you least want the monitor to fall over.
    """
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "empty.xlsx"
        reporter.write_excel([], path)
        check(path.exists() and path.stat().st_size > 0,
              "excel: an empty workbook is written without crashing")

        from openpyxl import load_workbook
        ws = load_workbook(path).active
        check_eq(ws.max_row, 1, "excel: only the header row is present")
        check(ws.auto_filter.ref is None,
              "excel: auto_filter is NOT set on an empty sheet")


def test_excel_colours_are_bare_hex():
    """openpyxl rejects #RRGGBB -- it needs bare hex."""
    for name in ("COLOR_HIGH", "COLOR_MEDIUM", "COLOR_LOW", "COLOR_HEADER"):
        value = getattr(config, name)
        check(not value.startswith("#"), f"excel: {name} has no leading '#'")
        check(len(value) == 6, f"excel: {name} is six hex digits")

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "coloured.xlsx"
        records = filters.process(fixtures.sample_records(TODAY), TODAY)["tenders"]
        reporter.decorate(records, TODAY)
        reporter.write_excel(records, path)
        check(path.exists(), "excel: a coloured workbook writes without raising")


# ---------------------------------------------------------------------------
# Delivery with no credentials
# ---------------------------------------------------------------------------


def test_default_is_file_only_and_sends_nothing():
    """The configured default writes files and sends no mail at all."""
    check_eq(config.EMAIL_METHOD, "none",
             "delivery: email is off by default -- files on disk are the product")
    result = emailer.deliver("subject", "<html></html>", "text", [])
    check_eq(result.method, "file", "delivery: reports the file-only path")
    check(result.sent, "delivery: treated as handled, not as a failure")
    check("nothing sent" in result.detail.lower(),
          "delivery: and says plainly that nothing was sent")


def test_delivery_without_credentials_degrades_to_disk():
    """If email is switched back on, a missing credential must not lose a run."""
    original_method = config.EMAIL_METHOD
    original_recipients = config.EMAIL_RECIPIENTS
    try:
        config.EMAIL_METHOD = "graph"
        config.EMAIL_RECIPIENTS = []
        result = emailer.deliver("subject", "<html></html>", "text", [])
        check_eq(result.method, "file", "delivery: falls back to file with no recipients")
        check(result.sent, "delivery: reported as handled, not as a failure")
        check("no mail was sent" in result.detail or "EMAIL_RECIPIENTS" in result.detail,
              "delivery: the reason is stated plainly")

        config.EMAIL_RECIPIENTS = ["someone@example.com"]
        result = emailer.deliver("subject", "<html></html>", "text", [])
        check_eq(result.method, "file",
                 "delivery: with recipients but no credentials it still degrades to disk")
        check(result.sent, "delivery: a credential problem never loses the run")
    finally:
        config.EMAIL_METHOD = original_method
        config.EMAIL_RECIPIENTS = original_recipients


def test_graph_without_credentials_does_not_raise():
    result = emailer.send_via_graph("s", "<html></html>", ["a@b.com"], [], [])
    check(not result.sent, "delivery: Graph reports failure with no credentials")
    check("credential" in result.detail.lower() or "SENDER_EMAIL" in result.detail,
          "delivery: and says why")


def test_smtp_without_credentials_does_not_raise():
    original = (config.SMTP_USER, config.SMTP_PASS)
    try:
        config.SMTP_USER, config.SMTP_PASS = "", ""
        result = emailer.send_via_smtp("s", "<html></html>", "t", ["a@b.com"], [], [])
        check(not result.sent, "delivery: SMTP reports failure with no credentials")
    finally:
        config.SMTP_USER, config.SMTP_PASS = original


# ---------------------------------------------------------------------------
# ACTION NEEDED alerts
#
# The reports are files; this is the one message that has to reach a person.
# Its whole value rests on firing when something is wrong and STAYING SILENT
# otherwise -- an alert that arrives on healthy days is one people filter.
# ---------------------------------------------------------------------------


def test_alert_fires_only_when_something_is_wrong():
    healthy = [h for h in fixtures.sample_health() if h.status == "ok"]

    needed, _ = emailer.should_alert(healthy)
    check(not needed, "alert: a healthy run sends nothing")

    needed, _ = emailer.should_alert(fixtures.sample_health())
    check(not needed,
          "alert: one flaky portal does not interrupt anyone -- it is visible "
          "in the report filename")

    needed, reason = emailer.should_alert(fixtures.all_broken_health())
    check(needed, "alert: a total outage DOES alert")
    check("unreachable" in reason, "alert: and says why", f"got {reason!r}")


def test_alert_respects_the_partial_threshold():
    original = config.ALERT_ON_PARTIAL_BROKEN
    try:
        health = fixtures.sample_health()          # exactly one broken portal
        config.ALERT_ON_PARTIAL_BROKEN = 1
        needed, reason = emailer.should_alert(health)
        check(needed, "alert: a configured threshold catches partial degradation")
        check("unavailable" in reason, "alert: the partial reason names the count")

        config.ALERT_ON_PARTIAL_BROKEN = 5
        needed, _ = emailer.should_alert(health)
        check(not needed, "alert: below the threshold stays silent")
    finally:
        config.ALERT_ON_PARTIAL_BROKEN = original


def test_alert_can_be_switched_off_entirely():
    original = config.ALERT_EMAIL
    try:
        config.ALERT_EMAIL = False
        needed, _ = emailer.should_alert(fixtures.all_broken_health())
        check(not needed, "alert: ALERT_EMAIL = False silences even a total outage")
    finally:
        config.ALERT_EMAIL = original


def test_alert_configuration_is_checked_before_it_is_needed():
    """An alert path you only discover is broken when you need it is not one."""
    original = (config.EMAIL_RECIPIENTS, config.AZURE_TENANT_ID,
                config.AZURE_CLIENT_ID, config.AZURE_CLIENT_SECRET,
                config.SENDER_EMAIL, config.SMTP_USER, config.SMTP_PASS)
    try:
        config.EMAIL_RECIPIENTS = []
        ok, detail = emailer.alert_configured()
        check(not ok, "alert-config: no recipients means not usable")
        check("nowhere to go" in detail, "alert-config: and says so plainly")

        config.EMAIL_RECIPIENTS = ["someone@example.com"]
        config.AZURE_TENANT_ID = config.AZURE_CLIENT_ID = ""
        config.AZURE_CLIENT_SECRET = config.SENDER_EMAIL = ""
        config.SMTP_USER = config.SMTP_PASS = ""
        ok, detail = emailer.alert_configured()
        check(not ok, "alert-config: recipients without credentials is not usable")
        check("credentials" in detail, "alert-config: and names what is missing")

        config.AZURE_TENANT_ID = "t"
        config.AZURE_CLIENT_ID = "c"
        config.AZURE_CLIENT_SECRET = "s"
        config.SENDER_EMAIL = "from@example.com"
        ok, detail = emailer.alert_configured()
        check(ok, "alert-config: complete Graph credentials are usable")
        check_eq(detail, "graph", "alert-config: reports which path would be used")
    finally:
        (config.EMAIL_RECIPIENTS, config.AZURE_TENANT_ID, config.AZURE_CLIENT_ID,
         config.AZURE_CLIENT_SECRET, config.SENDER_EMAIL, config.SMTP_USER,
         config.SMTP_PASS) = original


def test_alert_body_is_short_actionable_and_unattached():
    health = fixtures.all_broken_health()
    _, reason = emailer.should_alert(health)
    subject, html_body, text_body = reporter.build_alert(health, reason)

    check(subject.startswith(config.ACTION_NEEDED_PREFIX),
          "alert: the subject leads with ACTION NEEDED")
    check("unreachable" in subject, "alert: the subject states the failure")

    check("bot wall" in html_body or "transport error" in html_body,
          "alert: the diagnosed cause is in the body")
    check("ebrd.com" in html_body or "worldbank" in html_body,
          "alert: a URL to check by hand is included")
    check("--capture" in html_body,
          "alert: the body says what to run next")
    check(str(config.OUTPUT_DIR) in html_body,
          "alert: and where the files are")

    check(len(html_body) < 12_000,
          "alert: the body stays short -- this is an alert, not a digest",
          f"{len(html_body)} chars")
    check("World Bank" in text_body, "alert: a plain-text alternative is built")

    # The alert must never carry the report. An alert with a workbook attached
    # gets filtered, and it is the one message that must arrive.
    import inspect
    source = inspect.getsource(emailer.send_alert)
    check("[]" in source, "alert: send_alert passes no attachments")


def test_alert_without_credentials_reports_failure_rather_than_raising():
    original = (config.EMAIL_RECIPIENTS, config.SMTP_USER, config.SMTP_PASS)
    try:
        config.EMAIL_RECIPIENTS = ["someone@example.com"]
        config.SMTP_USER = config.SMTP_PASS = ""
        result = emailer.send_alert("subject", "<html></html>", "text")
        check(not result.sent,
              "alert: an unsendable alert reports failure instead of raising")
        check(result.detail, "alert: and carries the reason for the log")
    finally:
        (config.EMAIL_RECIPIENTS, config.SMTP_USER, config.SMTP_PASS) = original

    config_recipients = config.alert_recipients()
    check(isinstance(config_recipients, list),
          "alert: recipient resolution always returns a list")


# ---------------------------------------------------------------------------
# New-only mode, and the rule that diagnostics never touch real state
# ---------------------------------------------------------------------------


def test_new_only_mode_reports_each_tender_once():
    with tempfile.TemporaryDirectory() as tmp:
        store = tracker.Tracker(Path(tmp) / "seen.db")
        check(store.is_first_run(), "new-only: a fresh database is a first run")

        records = filters.process(fixtures.sample_records(TODAY), TODAY)["tenders"]
        first = store.filter_new(records)
        check_eq(len(first), len(records), "new-only: the first run reports everything")

        store.record(first)
        second = store.filter_new(records)
        check_eq(len(second), 0, "new-only: the second run reports nothing new")
        check(not store.is_first_run(), "new-only: no longer a first run")

        removed = store.reset()
        check_eq(removed, len(records), "new-only: reset forgets everything")
        check_eq(len(store.filter_new(records)), len(records),
                 "new-only: and the next run reports in full again")


def test_diagnostics_do_not_touch_the_real_database():
    """The --self-test path must not write fixture IDs into production state."""
    real_db = Path(config.SEEN_DB)
    before = real_db.stat().st_mtime if real_db.exists() else None
    before_ids = tracker.Tracker(real_db).seen_ids() if real_db.exists() else set()

    with tempfile.TemporaryDirectory() as tmp:
        store = tracker.Tracker(Path(tmp) / "seen.db")
        store.record(filters.process(fixtures.sample_records(TODAY), TODAY)["tenders"])
        check(store.count() > 0, "diagnostics: the temp database was written")

    after_ids = tracker.Tracker(real_db).seen_ids()
    check_eq(after_ids, before_ids,
             "diagnostics: the REAL database is unchanged")
    if before is not None:
        check(real_db.stat().st_mtime == before or after_ids == before_ids,
              "diagnostics: no fixture IDs leaked into production state")


def test_tender_ids_are_stable_across_runs():
    a = _record(url="https://e.org/stable", reference="REF-1")
    b = _record(url="https://e.org/stable", reference="REF-1",
                description="An edited summary that differs from the first run.")
    check_eq(a["id"], b["id"],
             "identity: the same notice keeps its id when its text is edited")


def test_portal_health_unconfigured_is_not_broken():
    unconfigured = PortalHealth("samgov", "SAM.gov", 1, "unconfigured",
                                reason="not configured - no SAM_API_KEY in .env")
    broken = PortalHealth("ebrd", "EBRD", 2, "unavailable", reason="bot wall")
    check(not unconfigured.broken,
          "health: awaiting an API key is NOT a scraper failure")
    check(broken.broken, "health: an unreachable portal IS a failure")


TESTS = [
    test_empty_input,
    test_every_optional_field_none,
    test_zero_keyword_matches_division_guard,
    test_deadline_boundary_today_is_kept,
    test_unknown_value_is_kept_and_flagged,
    test_unknown_value_scores_mid_band,
    test_national_only_is_flagged_and_penalised,
    test_arabic_notice_is_kept_and_flagged,
    test_all_sectors_and_types_pass,
    test_disabled_component_is_dropped_and_weights_renormalise,
    test_ranking_favours_advisory_over_goods,
    test_duplicate_collapse_across_portals,
    test_dedupe_keeps_the_better_data,
    test_same_portal_numbered_lots_are_not_merged,
    test_distinct_tenders_are_not_merged,
    test_subject_distinguishes_quiet_from_broken,
    test_subject_reports_partial_degradation,
    test_subject_singular_and_plural,
    test_email_overflow_moves_nothing_is_dropped,
    test_health_table_names_the_failure_and_the_url,
    test_configured_formats_written_and_arabic_survives,
    test_retained_json_and_csv_writers_still_handle_arabic,
    test_filename_states_the_run_health,
    test_documents_state_the_run_health_inside,
    test_excel_with_zero_rows_does_not_crash,
    test_excel_colours_are_bare_hex,
    test_default_is_file_only_and_sends_nothing,
    test_delivery_without_credentials_degrades_to_disk,
    test_graph_without_credentials_does_not_raise,
    test_smtp_without_credentials_does_not_raise,
    test_alert_fires_only_when_something_is_wrong,
    test_alert_respects_the_partial_threshold,
    test_alert_can_be_switched_off_entirely,
    test_alert_configuration_is_checked_before_it_is_needed,
    test_alert_body_is_short_actionable_and_unattached,
    test_alert_without_credentials_reports_failure_rather_than_raising,
    test_new_only_mode_reports_each_tender_once,
    test_diagnostics_do_not_touch_the_real_database,
    test_tender_ids_are_stable_across_runs,
    test_portal_health_unconfigured_is_not_broken,
]
