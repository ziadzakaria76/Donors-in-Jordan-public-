"""
Reporter agent: the email body, the subject line, and five output files.

Two things here are load-bearing.

THE SUBJECT LINE CARRIES PORTAL HEALTH (Q15). If the subject says
"0 opportunities" whether every portal failed or every portal worked and
nothing matched, a dead monitor goes unnoticed for weeks and you find out when
someone asks why you missed a bid. So:

    Jordan Tenders - 7 new opportunities
    Jordan Tenders - no new opportunities (13/13 portals OK)
    Jordan Tenders - 4 new, 3 portals unavailable
    ACTION NEEDED: Jordan Tenders - all 13 portals unreachable

THE EMAIL NEVER DROPS A TENDER. Outlook clips messages over roughly 100 KB and
hides the tail without saying so, so full detail is rendered for the top
MAX_INLINE_TENDERS by score and the remainder MOVES to a compact table with a
pointer to the attachments. Moving is not dropping.
"""

from __future__ import annotations

import csv
import html
import json
import logging
from datetime import date, datetime
from pathlib import Path

from .. import config
from ..utils import money
from ..utils.dates import days_until, fmt
from ..utils.text import truncate

log = logging.getLogger(__name__)

COLUMNS = [
    ("score", "Score"),
    ("title", "Title"),
    ("portal_name", "Portal"),
    ("notice_type", "Notice type"),
    ("sector", "Sector"),
    ("posted_date", "Published"),
    ("closing_date", "Deadline"),
    ("days_left", "Days left"),
    ("value_display", "Estimated value"),
    ("language", "Language"),
    ("flags_display", "Flags"),
    ("eligibility", "Eligibility"),
    ("contact", "Contact"),
    ("url", "Link"),
]


# ---------------------------------------------------------------------------
# Presentation helpers
# ---------------------------------------------------------------------------


def decorate(tenders: list[dict], today: date | None = None) -> list[dict]:
    """Add display-only fields. Never mutates the underlying data meaning."""
    today = today or date.today()
    for t in tenders:
        t["portal_name"] = config.PORTAL_NAMES.get(t.get("portal"), t.get("portal"))
        t["value_display"] = money.format_usd(t.get("estimated_value_usd"))
        left = days_until(t.get("closing_date"), today)
        t["days_left"] = "" if left is None else left
        t["flags_display"] = "; ".join(t.get("flags") or [])
        t["posted_display"] = fmt(t.get("posted_date"))
        t["closing_display"] = fmt(t.get("closing_date"))
    return tenders


def build_subject(reported: int, health: list, first_run: bool = False) -> str:
    """Subject line that states the run's health, not just its count."""
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    ok = total - len(broken)
    prefix = config.EMAIL_SUBJECT_PREFIX

    if total and len(broken) == total:
        return f"{config.ACTION_NEEDED_PREFIX}: {prefix} - all {total} portals unreachable"

    first = " (first run - full pipeline)" if first_run and reported else ""

    if broken:
        noun = "opportunity" if reported == 1 else "opportunities"
        return (f"{prefix} - {reported} new {noun}, "
                f"{len(broken)} of {total} portals unavailable{first}")

    if reported == 0:
        return f"{prefix} - no new opportunities ({ok}/{total} portals OK)"

    noun = "opportunity" if reported == 1 else "opportunities"
    return f"{prefix} - {reported} new {noun}{first}"


# ---------------------------------------------------------------------------
# HTML email body
# ---------------------------------------------------------------------------

_CSS = """
body{font-family:Segoe UI,Helvetica,Arial,sans-serif;font-size:14px;color:#222;line-height:1.45}
h1{font-size:19px;margin:0 0 4px}h2{font-size:16px;margin:22px 0 8px;
   border-bottom:2px solid #1F4E79;padding-bottom:3px;color:#1F4E79}
table{border-collapse:collapse;width:100%;margin:8px 0;font-size:13px}
th{background:#1F4E79;color:#fff;text-align:left;padding:6px;font-weight:600}
td{border-bottom:1px solid #ddd;padding:6px;vertical-align:top}
.t{border:1px solid #ddd;border-left:5px solid #1F4E79;padding:10px 12px;margin:10px 0;
   border-radius:3px;background:#fafafa}
.t h3{margin:0 0 6px;font-size:15px}
.hi{border-left-color:#2E7D32}.mid{border-left-color:#F9A825}.lo{border-left-color:#C62828}
.meta{font-size:12px;color:#555;margin:2px 0}
.flag{display:inline-block;background:#FFF3CD;border:1px solid #FFE08A;color:#7A5B00;
      padding:1px 6px;border-radius:9px;font-size:11px;margin:2px 3px 2px 0}
.err{background:#FDECEA;border:1px solid #F5C2C0;padding:10px;border-radius:3px}
.rtl{direction:rtl;text-align:right}
.small{font-size:12px;color:#666}
"""


def _band(score: float) -> str:
    return "hi" if score >= 70 else ("mid" if score >= 40 else "lo")


def _e(value) -> str:
    return html.escape(str(value if value is not None else ""))


def _tender_block(t: dict) -> str:
    rtl = ' class="rtl"' if t.get("language") == "ar" else ""
    flags = "".join(f'<span class="flag">{_e(f)}</span>' for f in (t.get("flags") or []))
    also = ""
    if t.get("also_on"):
        also = f'<div class="meta">Also published on: {_e(", ".join(t["also_on"]))}</div>'
    for extra in (t.get("also_urls") or []):
        # The same notice re-published under a second id. Both pages are real,
        # so the merge shows both rather than picking one and losing the other.
        also += (f'<div class="meta">Also published as: '
                 f'<a href="{_e(extra)}">{_e(extra)}</a></div>')
    link = (f'<div class="meta"><a href="{_e(t.get("url"))}">Open the notice</a></div>'
            if t.get("url") else '<div class="meta">No direct link published</div>')
    desc = truncate(t.get("description") or "", config.DESCRIPTION_CHAR_LIMIT)
    contact = (f'<div class="meta"><b>Contact:</b> {_e(t["contact"])}</div>'
               if t.get("contact") else "")
    days = t.get("days_left")
    days_txt = f" ({days} days left)" if days != "" and days is not None else ""

    return f"""
<div class="t {_band(t.get('score', 0))}"{rtl}>
  <h3>{_e(t.get('title'))}</h3>
  <div class="meta"><b>Score {t.get('score')}</b> &middot; {_e(t.get('portal_name'))}
      &middot; {_e(t.get('sector'))} &middot; {_e(t.get('notice_type') or 'Type unspecified')}</div>
  <div class="meta"><b>Published:</b> {_e(t.get('posted_display'))}
      &nbsp;|&nbsp; <b>Deadline:</b> {_e(t.get('closing_display'))}{_e(days_txt)}
      &nbsp;|&nbsp; <b>Value:</b> {_e(t.get('value_display'))}</div>
  {contact}{also}
  <div class="meta">{_e(desc)}</div>
  {flags}
  {link}
</div>"""


def _health_table(health: list) -> str:
    rows = []
    for h in health:
        if h.status == "ok":
            status = '<span style="color:#2E7D32">OK</span>'
            detail = f"{h.count} Jordan notice(s)"
            if h.scanned is not None and h.scanned != h.count:
                detail += f" &middot; {h.scanned} read, {h.scanned - h.count} not Jordan"
            if h.layer:
                detail += f" &middot; via {_e(h.layer)} layer (quality {h.quality:.2f})"
        elif h.status == "unconfigured":
            status = '<span style="color:#8A6D00">NOT CONFIGURED</span>'
            detail = _e(h.reason)
        elif h.status == "no listing":
            # Not red. This source publishes nothing to read, which is a fact
            # about the source rather than a fault in the monitor.
            status = '<span style="color:#8A6D00">NO LISTING</span>'
            detail = _e(h.reason)
        else:
            status = '<span style="color:#C62828"><b>UNAVAILABLE</b></span>'
            urls = "<br>".join(f'<a href="{_e(u)}">{_e(u)}</a>' for u in h.urls[:2])
            detail = f"{_e(h.reason)}<br><span class='small'>Check by hand: {urls}</span>"
        tier = config.TIER_LABELS.get(h.tier, "")
        rows.append(f"<tr><td>{_e(h.name)}<br><span class='small'>{_e(tier)}</span></td>"
                    f"<td>{status}</td><td>{detail}</td></tr>")
    return ("<table><tr><th>Portal</th><th>Status</th><th>Detail</th></tr>"
            + "".join(rows) + "</table>")


def _overflow_table(tenders: list[dict]) -> str:
    rows = []
    for t in tenders:
        title = (f'<a href="{_e(t.get("url"))}">{_e(t.get("title"))}</a>'
                 if t.get("url") else _e(t.get("title")))
        rows.append(
            f"<tr><td>{_e(t.get('score'))}</td><td>{title}</td>"
            f"<td>{_e(t.get('portal_name'))}</td><td>{_e(t.get('closing_display'))}</td>"
            f"<td>{_e(t.get('value_display'))}</td></tr>")
    return ("<table><tr><th>Score</th><th>Title</th><th>Portal</th>"
            "<th>Deadline</th><th>Value</th></tr>" + "".join(rows) + "</table>")


def build_email_html(result: dict, health: list, first_run: bool = False) -> str:
    tenders = result["tenders"]
    today = date.today()
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if h.status not in ("unconfigured", "no listing")])

    banner = ""
    if broken and len(broken) == total:
        banner = (f'<div class="err"><b>ACTION NEEDED.</b> All {total} portals were '
                  f'unreachable this run. This report is empty because nothing could '
                  f'be read, not because nothing was published. See the portal status '
                  f'table below.</div>')
    elif broken:
        names = ", ".join(h.name for h in broken)
        banner = (f'<div class="err"><b>{len(broken)} of {total} portals unavailable:</b> '
                  f'{_e(names)}. The opportunities below are therefore a partial '
                  f'picture. Details in the portal status table.</div>')

    first_note = ""
    if first_run and tenders:
        first_note = ('<p class="small">This is the first run, so the seen-tenders '
                      'database was empty and the whole open pipeline is reported. '
                      'From the next run onward only new notices appear.</p>')

    inline = tenders[:config.MAX_INLINE_TENDERS]
    overflow = tenders[config.MAX_INLINE_TENDERS:]

    parts = [
        f"<html><head><meta charset='utf-8'><style>{_CSS}</style></head><body>",
        "<h1>Jordan Tender Intelligence</h1>",
        f"<div class='small'>{today.strftime('%A %d %B %Y')} &middot; "
        f"{len(tenders)} opportunit{'y' if len(tenders)==1 else 'ies'} reported "
        f"&middot; {result['scanned']} notices scanned across {total} portals</div>",
        banner, first_note,
    ]

    if tenders:
        parts.append(f"<h2>Opportunities ({len(inline)} shown in full)</h2>")
        parts.extend(_tender_block(t) for t in inline)
        if overflow:
            parts.append(
                f"<h2>Further opportunities ({len(overflow)})</h2>"
                f"<p class='small'>Listed compactly to keep this email under Outlook's "
                f"clipping limit. Nothing has been dropped &mdash; full detail for every "
                f"one of these is in the attached Word and Excel files.</p>")
            parts.append(_overflow_table(overflow))
    else:
        if not broken:
            parts.append("<h2>No new opportunities</h2><p>Every portal was read "
                         "successfully and nothing new matched. This is a genuine "
                         "quiet run, not a failure &mdash; see the status table.</p>")

    dropped = result.get("dropped") or {}
    if dropped:
        items = ", ".join(f"{v} {k}" for k, v in sorted(dropped.items()))
        parts.append(f"<p class='small'><b>Filtered out this run:</b> {_e(items)}. "
                     f"Merged as duplicates across portals: "
                     f"{result.get('merged_duplicates', 0)}.</p>")

    parts.append("<h2>Portal status</h2>")
    parts.append(_health_table(health))

    weights = ", ".join(f"{k} {v:.1f}" for k, v in result.get("weights", {}).items())
    parts.append(f"<p class='small'>Scoring weights in force (renormalised to 100 "
                 f"after dropping components with a disabled filter): {_e(weights)}.</p>")
    parts.append("</body></html>")
    return "\n".join(p for p in parts if p)


def build_text_body(result: dict, health: list) -> str:
    """Plain-text alternative, for clients that will not render HTML."""
    lines = [f"Jordan Tender Intelligence - {date.today().isoformat()}", ""]
    for t in result["tenders"]:
        lines += [
            f"[{t.get('score')}] {t.get('title')}",
            f"    {t.get('portal_name')} | {t.get('sector')} | "
            f"deadline {t.get('closing_display')} | {t.get('value_display')}",
        ]
        if t.get("flags"):
            lines.append(f"    flags: {'; '.join(t['flags'])}")
        if t.get("url"):
            lines.append(f"    {t['url']}")
        lines.append("")
    lines.append("Portal status:")
    for h in health:
        lines.append(f"  {h.name}: {h.status.upper()}"
                     + (f" ({h.count})" if h.ok else f" - {h.reason}"))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Output files
# ---------------------------------------------------------------------------


def _stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M")


_SLUG_UNSAFE = str.maketrans({c: "-" for c in r'\/:*?"<>| '})


def run_slug(reported: int, health: list) -> str:
    """Filename fragment stating the run's health, not just its count.

    This carries what the email subject line used to. With output going to disk
    instead of an inbox, a run where every portal failed and a run where nothing
    new was published both leave a file behind -- and identically-named files
    make a dead monitor invisible. The folder listing must show the difference:

        2026-08-03_7-opportunities.docx
        2026-08-03_no-new-opportunities_13-of-13-portals-OK.docx
        2026-08-03_ACTION-NEEDED_all-13-portals-unreachable.docx
    """
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    ok = total - len(broken)

    if total and len(broken) == total:
        slug = f"{config.ACTION_NEEDED_PREFIX}-all-{total}-portals-unreachable"
    elif broken:
        noun = "opportunity" if reported == 1 else "opportunities"
        slug = f"{reported}-{noun}-{len(broken)}-of-{total}-portals-unavailable"
    elif reported == 0:
        slug = f"no-new-opportunities-{ok}-of-{total}-portals-OK"
    else:
        noun = "opportunity" if reported == 1 else "opportunities"
        slug = f"{reported}-{noun}"

    return slug.translate(_SLUG_UNSAFE).replace("--", "-").strip("-")


def status_line(reported: int, health: list) -> str:
    """One-sentence run status, shown at the top of both documents."""
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])

    if total and len(broken) == total:
        return (f"{config.ACTION_NEEDED_PREFIX}: all {total} portals were unreachable. "
                f"This report is empty because nothing could be read, NOT because "
                f"nothing was published.")
    if broken:
        return (f"{reported} opportunities found. {len(broken)} of {total} portals "
                f"were unavailable, so this is a partial picture: "
                + ", ".join(h.name for h in broken))
    if reported == 0:
        return (f"No new opportunities. All {total} portals were read successfully "
                f"and nothing new matched -- a genuine quiet run, not a failure.")
    return f"{reported} new opportunities found. All {total} portals were read successfully."


def _cell(t: dict, key: str):
    value = t.get(key)
    if isinstance(value, date):
        return value.isoformat()
    return "" if value is None else value


def write_excel(tenders: list[dict], path: Path, health: list | None = None,
                reported: int | None = None) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Jordan tenders"

    headers = [label for _, label in COLUMNS]
    ws.append(headers)

    # openpyxl requires BARE hex. "#1F4E79" raises; "1F4E79" is correct.
    header_fill = PatternFill(start_color=config.COLOR_HEADER,
                              end_color=config.COLOR_HEADER, fill_type="solid")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(vertical="center", wrap_text=True)

    for t in tenders:
        ws.append([_cell(t, key) for key, _ in COLUMNS])
        fill_hex = (config.COLOR_HIGH if (t.get("score") or 0) >= 70
                    else config.COLOR_MEDIUM if (t.get("score") or 0) >= 40
                    else config.COLOR_LOW)
        fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
        for cell in ws[ws.max_row]:
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = [8, 60, 24, 16, 22, 12, 12, 10, 18, 10, 34, 34, 28, 46]
    for i, width in enumerate(widths[:len(COLUMNS)], start=1):
        ws.column_dimensions[get_column_letter(i)].width = width
    ws.freeze_panes = "A2"

    # Setting auto_filter over a header-only range is a classic crash, and it
    # happens on precisely the quiet day when nothing matched.
    if tenders:
        ws.auto_filter.ref = f"A1:{get_column_letter(len(COLUMNS))}{ws.max_row}"

    # Portal health goes on its own sheet rather than above the headers, which
    # would shift every column and break the auto_filter range. Without it, an
    # empty workbook from a total outage is indistinguishable from an empty
    # workbook from a quiet week.
    if health is not None:
        _append_status_sheet(wb, health, reported if reported is not None else len(tenders))

    wb.save(path)
    return path


def _append_status_sheet(wb, health: list, reported: int) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    ws = wb.create_sheet("Run status")
    ws["A1"] = "Run status"
    ws["A1"].font = Font(bold=True, size=13)
    ws["A2"] = status_line(reported, health)
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws["A3"] = f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}"

    header_row = 5
    for col, label in enumerate(("Portal", "Tier", "Status", "Notices", "Detail",
                                 "URL to check by hand"), start=1):
        cell = ws.cell(row=header_row, column=col, value=label)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = PatternFill(start_color=config.COLOR_HEADER,
                                end_color=config.COLOR_HEADER, fill_type="solid")

    for offset, h in enumerate(health, start=1):
        row = header_row + offset
        ws.cell(row=row, column=1, value=h.name)
        ws.cell(row=row, column=2, value=config.TIER_LABELS.get(h.tier, ""))
        ws.cell(row=row, column=3, value=h.status.upper())
        ws.cell(row=row, column=4, value=h.count if h.ok else "")
        ws.cell(row=row, column=5, value=h.reason or (
            f"read via {h.layer} layer" if h.layer else ""))
        ws.cell(row=row, column=6, value=h.urls[0] if h.urls and not h.ok else "")

        fill_hex = (config.COLOR_HIGH if h.ok
                    else config.COLOR_MEDIUM if h.status in ("unconfigured", "no listing")
                    else config.COLOR_LOW)
        fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
        for col in range(1, 7):
            ws.cell(row=row, column=col).fill = fill
            ws.cell(row=row, column=col).alignment = Alignment(vertical="top",
                                                               wrap_text=True)

    for i, width in enumerate((46, 20, 16, 10, 60, 60), start=1):
        ws.column_dimensions[get_column_letter(i)].width = width


def write_docx(tenders: list[dict], health: list, result: dict, path: Path) -> Path:
    """Word bid-review pack -- the artefact that circulates and takes comments."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Jordan Tender Intelligence", level=0)
    subtitle = doc.add_paragraph(date.today().strftime("%A %d %B %Y"))
    subtitle.runs[0].font.size = Pt(11)
    doc.add_paragraph(
        f"{len(tenders)} opportunities reported from {result.get('scanned', 0)} "
        f"notices scanned. Ranked by fit score."
    )

    # Run status on the first page, in words. Someone opening this file weeks
    # later should not have to infer from an empty list whether the pipeline was
    # quiet or the scrapers were dead.
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    status = doc.add_paragraph()
    run = status.add_run(status_line(len(tenders), health))
    run.bold = True
    if broken:
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    if not tenders and not broken:
        doc.add_paragraph(
            f"All {total} portals were read successfully and returned no new "
            f"Jordan opportunities. This file is evidence the monitor ran, not "
            f"evidence that it failed.")

    for t in tenders:
        doc.add_heading(t.get("title") or "(untitled)", level=2)
        meta = doc.add_paragraph()
        meta.add_run(f"Score {t.get('score')}").bold = True
        meta.add_run(f"  |  {t.get('portal_name')}  |  {t.get('sector')}"
                     f"  |  {t.get('notice_type') or 'Type unspecified'}")
        doc.add_paragraph(
            f"Published: {t.get('posted_display')}    "
            f"Deadline: {t.get('closing_display')}    "
            f"Estimated value: {t.get('value_display')}")
        if t.get("flags"):
            flagline = doc.add_paragraph()
            flagline.add_run("Flags: " + "; ".join(t["flags"])).italic = True
        if t.get("contact"):
            doc.add_paragraph(f"Contact: {t['contact']}")
        if t.get("also_on"):
            doc.add_paragraph(f"Also published on: {', '.join(t['also_on'])}")
        for extra in (t.get("also_urls") or []):
            doc.add_paragraph(f"Also published as: {extra}")
        if t.get("description"):
            body = doc.add_paragraph(truncate(t["description"], 2000))
            if t.get("language") == "ar":
                body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if t.get("url"):
            doc.add_paragraph(t["url"])

    doc.add_page_break()
    doc.add_heading("Portal status", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Portal", "Status", "Detail"
    for h in health:
        cells = table.add_row().cells
        cells[0].text = h.name
        cells[1].text = h.status.upper()
        cells[2].text = (f"{h.count} notice(s)" if h.ok
                         else f"{h.reason}  ({'; '.join(h.urls[:1])})")

    doc.save(path)
    return path


# The report schema the Android app reads. Bump it when a field changes
# meaning or disappears -- never when one is added, since an older app ignores
# what it does not know. The app refuses to render a schema it was not written
# for, which is a worse experience than guessing exactly once and a better one
# than a screen of half-parsed nonsense every day after that.
REPORT_SCHEMA = 1


def run_status(reported: int, health: list) -> str:
    """The run's outcome as one machine-readable word.

    The same four cases the filename slug and the status line already
    distinguish, named rather than re-derived: an app that re-implemented
    "which of these is a bad run" would eventually disagree with the documents,
    and the disagreement would be invisible.
    """
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health
                 if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    if total and len(broken) == total:
        return "action_needed"
    if broken:
        return "partial"
    if reported == 0:
        return "quiet"
    return "ok"


# Display-only fields whose "unknown" value is an empty string, because that
# is what belongs in a Word table. JSON is typed and a client is not a table:
# "" would have to be special-cased at every call site on the phone, and the
# one that got missed would render as a real value.
_JSON_EMPTY_TO_NULL = ("days_left",)


def _json_tender(tender: dict) -> dict:
    out = {k: v for k, v in tender.items() if not k.startswith("_")}
    for key in _JSON_EMPTY_TO_NULL:
        if out.get(key) == "":
            out[key] = None
    return out


def write_json(tenders: list[dict], health: list, path: Path,
               result: dict | None = None) -> Path:
    """The whole report as structured data, for the Android app.

    NOT a deliverable and not a convenience dump: it is the app's only way to
    read a run. GitHub's REST API does not expose a job's step summary --
    summaries live in an internal container the artifacts API does not list --
    so the markdown rendered onto the run page is unreachable from a client.
    This file is uploaded as an artifact alongside the Word and Excel packs,
    and the app downloads and parses it.

    Every field the report shows on paper is here, including the ones that make
    a quiet run distinguishable from a broken one. In particular `scanned` is
    null when a portal never filters, and MUST NOT be rendered as 0: "read
    nothing" and "read 500 and none were Jordan" are the two diagnoses that
    number exists to separate.
    """
    def encode(obj):
        if isinstance(obj, date):
            return obj.isoformat()
        raise TypeError(type(obj))

    result = result or {}
    considered = [h for h in health
                  if getattr(h, "status", "") not in ("unconfigured", "no listing")]
    broken = [h for h in health if getattr(h, "broken", False)]

    payload = {
        "schema": REPORT_SCHEMA,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "run": {
            "status": run_status(len(tenders), health),
            "status_line": status_line(len(tenders), health),
            "slug": run_slug(len(tenders), health),
            "opportunity_count": len(tenders),
            # Notices read before filtering, across every portal. The
            # difference between this and opportunity_count is the pipeline
            # working, not notices going missing.
            "scanned": result.get("scanned", 0),
            "merged_duplicates": result.get("merged_duplicates", 0),
            "dropped": dict(result.get("dropped") or {}),
            "portals_total": len(considered),
            "portals_ok": len(considered) - len(broken),
            "portals_broken": len(broken),
            "new_only": config.NEW_ONLY_MODE,
        },
        "tender_count": len(tenders),
        "tenders": [_json_tender(t) for t in tenders],
        "portals": [
            {"key": h.key, "name": h.name, "tier": h.tier,
             "tier_label": config.TIER_LABELS.get(h.tier, ""),
             "status": h.status, "count": h.count,
             # null, not 0. See the docstring.
             "scanned": h.scanned,
             "reason": h.reason, "urls": list(h.urls),
             "layer": h.layer, "quality": h.quality}
            for h in health
        ],
    }
    # ensure_ascii=False so Arabic survives as Arabic rather than \uXXXX.
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False, default=encode),
                    encoding="utf-8")
    return path


def write_csv(tenders: list[dict], path: Path) -> Path:
    # utf-8-sig: without the BOM, Excel opens Arabic as mojibake.
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.writer(fh)
        writer.writerow([label for _, label in COLUMNS])
        for t in tenders:
            writer.writerow([_cell(t, key) for key, _ in COLUMNS])
    return path


def write_markdown(tenders: list[dict], health: list, result: dict, path: Path) -> Path:
    """Markdown summary, for the GitHub Actions run page.

    Not a deliverable -- the Word and Excel files are. This exists so results
    are readable on a phone without downloading a .docx: GitHub renders it into
    the workflow run page, so tapping "Run workflow" and then reading the
    outcome is two taps and no file manager.
    """
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    lines: list[str] = []

    if broken and len(broken) == total:
        lines.append(f"# {config.ACTION_NEEDED_PREFIX}")
    elif tenders:
        lines.append(f"# {len(tenders)} Jordan opportunit"
                     f"{'y' if len(tenders) == 1 else 'ies'}")
    else:
        lines.append("# No new opportunities")

    lines += [f"\n{status_line(len(tenders), health)}\n",
              f"_{result.get('scanned', 0)} notices scanned across {total} portals "
              f"on {date.today().isoformat()}._\n"]

    if tenders:
        lines += ["\n## Opportunities\n",
                  "| Score | Title | Portal | Deadline | Value |",
                  "|---:|---|---|---|---|"]
        for t in tenders[:40]:
            title = (t.get("title") or "").replace("|", "\\|")
            if t.get("url"):
                title = f"[{title}]({t['url']})"
            lines.append(
                f"| {t.get('score')} | {title} | {t.get('portal_name')} | "
                f"{t.get('closing_display')} | {t.get('value_display')} |")
        if len(tenders) > 40:
            lines.append(f"\n_{len(tenders) - 40} more in the attached Word and "
                         f"Excel files._")

    lines += ["\n## Portal status\n",
              "| Portal | Status | Detail |", "|---|---|---|"]
    for h in health:
        if h.status == "ok":
            detail = f"{h.count} Jordan notice(s)"
            if h.scanned is not None and h.scanned != h.count:
                detail += f" — {h.scanned} read, {h.scanned - h.count} not Jordan"
            if h.layer:
                detail += f" via `{h.layer}`"
        else:
            detail = (h.reason or "").replace("|", "\\|")[:160]
        mark = {"ok": "OK", "unconfigured": "not configured",
                "no listing": "no listing published"}.get(h.status, "**UNAVAILABLE**")
        lines.append(f"| {h.name} | {mark} | {detail} |")

    if broken:
        lines.append("\n### Check by hand\n")
        for h in broken:
            if h.urls:
                lines.append(f"- {h.name}: <{h.urls[0]}>")

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_html(body_html: str, path: Path) -> Path:
    path.write_text(body_html, encoding="utf-8")
    return path


def write_outputs(result: dict, health: list, body_html: str,
                  output_dir: Path | None = None) -> dict[str, Path]:
    """Write every configured output format. Returns format -> path."""
    output_dir = Path(output_dir or config.OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    tenders = result["tenders"]
    stamp = _stamp()
    written: dict[str, Path] = {}

    # The filename states the run's health, so the folder listing alone
    # distinguishes a quiet day from a dead monitor.
    if config.HEALTH_IN_FILENAME:
        base = f"jordan_tenders_{stamp}_{run_slug(len(tenders), health)}"
    else:
        base = f"jordan_tenders_{stamp}"

    for fmt_name in config.OUTPUT_FORMATS:
        try:
            if fmt_name == "excel":
                written["excel"] = write_excel(tenders, output_dir / f"{base}.xlsx",
                                               health, len(tenders))
            elif fmt_name == "docx":
                written["docx"] = write_docx(tenders, health, result,
                                             output_dir / f"{base}.docx")
            elif fmt_name == "json":
                # The app reads this one. It is given `result` as well, so the
                # run's own counts travel with it rather than being recomputed
                # on a handset from a list that has already been filtered.
                written["json"] = write_json(tenders, health,
                                             output_dir / f"{base}.json", result)
            elif fmt_name == "csv":
                written["csv"] = write_csv(tenders, output_dir / f"{base}.csv")
            elif fmt_name == "html":
                written["html"] = write_html(body_html, output_dir / f"{base}.html")
            elif fmt_name == "markdown":
                written["markdown"] = write_markdown(tenders, health, result,
                                                     output_dir / f"{base}.md")
        except Exception as exc:  # noqa: BLE001 - one format must not lose the rest
            log.error("could not write %s output: %s", fmt_name, exc)

    return written


def build_alert(health: list, reason: str, written: dict[str, Path] | None = None
                ) -> tuple[str, str, str]:
    """The ACTION NEEDED alert: subject, HTML body, plain-text body.

    Deliberately short and unattached. This is not the report -- it is the one
    message that has to reach a person, so it says what broke, what to check,
    and nothing else. An alert that arrives with a 2 MB workbook attached gets
    filtered; an alert that reads like a daily digest gets ignored.
    """
    broken = [h for h in health if getattr(h, "broken", False)]
    total = len([h for h in health if getattr(h, "status", "") not in ("unconfigured", "no listing")])
    today = date.today()

    subject = f"{config.ALERT_SUBJECT_PREFIX} - {reason} ({today.isoformat()})"

    rows = "".join(
        f"<tr><td>{_e(h.name)}</td><td>{_e(h.reason)}</td>"
        f"<td><a href=\"{_e(h.urls[0] if h.urls else '')}\">"
        f"{_e(h.urls[0] if h.urls else 'no URL recorded')}</a></td></tr>"
        for h in broken
    )

    where = ""
    if written:
        where = (f"<p class='small'>The run still wrote its files to "
                 f"<code>{_e(config.OUTPUT_DIR)}</code>, but they are empty or "
                 f"partial for the reason above.</p>")

    html = f"""<html><head><meta charset='utf-8'><style>{_CSS}</style></head><body>
<h1>Jordan Tender Monitor needs attention</h1>
<div class="err"><b>{_e(reason.capitalize())}</b> on {today.strftime('%A %d %B %Y')}.
This is not a quiet day &mdash; the monitor could not read
{"any of its sources" if len(broken) == total else "some of its sources"},
so any opportunity published today may have been missed.</div>

<h2>What failed</h2>
<table><tr><th>Portal</th><th>Diagnosed cause</th><th>Check by hand</th></tr>
{rows}</table>
{where}

<h2>What to do</h2>
<ol>
  <li>Open one of the URLs above in a browser <b>on the server</b>. If it does
      not load, the problem is the network or a firewall, not the scraper.</li>
  <li>If it loads fine by hand, the site has changed. Run
      <code>python run.py --capture PORTAL</code> for that portal &mdash; it
      reports which extraction layer works and the selectors the page now
      uses.</li>
  <li>A <b>bot wall</b> means the site is blocking automated access: try a
      different network, or install Playwright.</li>
</ol>
<p class="small">You are receiving this because the monitor alerts on failure
only. A run that works quietly sends nothing &mdash; the reports are files in
{_e(config.OUTPUT_DIR)}.</p>
</body></html>"""

    text_lines = [
        f"{config.ALERT_SUBJECT_PREFIX} - {reason}",
        f"{today.isoformat()}",
        "",
        "Failed portals:",
    ]
    for h in broken:
        text_lines.append(f"  {h.name}: {h.reason}")
        if h.urls:
            text_lines.append(f"    check: {h.urls[0]}")
    text_lines += ["", f"Reports (empty or partial) are in {config.OUTPUT_DIR}"]

    return subject, html, "\n".join(text_lines)


def attachments_for_email(written: dict[str, Path]) -> list[Path]:
    return [written[f] for f in config.EMAIL_ATTACH_FORMATS if f in written]
